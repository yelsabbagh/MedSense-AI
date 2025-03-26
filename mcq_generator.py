# mcq_generator.py

import google.generativeai as genai
import config
import re
import os
import time
import pandas as pd
from google.api_core.exceptions import ResourceExhausted
import pypandoc
from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import lxml.etree as etree
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate # For rendering context
import traceback

# --- System Instructions ---
SYSTEM_INSTRUCTION_GENERATOR = """You are a medical exam question generator. Your task is to create Multiple-Choice Questions (MCQs) from provided lecture material, strictly following the guidelines in the attached `rules.txt` file. Focus on clinical reasoning, accuracy, and adherence to medical exam standards (e.g., USMLE)."""
SYSTEM_INSTRUCTION_VERIFIER = """You are a medical exam question verifier and corrector. Your task is to analyze Multiple-Choice Questions (MCQs) for any violations of the rules provided, correct them if necessary, and ensure they adhere to the specified output format (Question, Options a-e, Correct Answer letter)."""

# --- Gemini API Call ---
def generate_with_retry(prompt, system_instruction, retries=5, delay=5):
    """Retries the generation request with system instruction, handling rate limits."""
    genai.configure(api_key=config.API_KEY)
    model = genai.GenerativeModel(
        model_name=config.GEMINI_MODEL,
        generation_config=config.generation_config,
        safety_settings=config.safety_settings,
        system_instruction=system_instruction
    )
    for attempt in range(retries):
        try:
            # print(f"    Attempting Gemini API call ({attempt + 1}/{retries})...")
            response = model.generate_content(prompt)
            if response and hasattr(response, 'text') and response.text.strip():
                # print(f"    Gemini call successful.")
                return response.text
            else:
                 # Handle cases where the response might be blocked or empty
                 block_reason = ""
                 if response and response.prompt_feedback and response.prompt_feedback.block_reason:
                     block_reason = f" (Block Reason: {response.prompt_feedback.block_reason})"
                 print(f"    Warning: Gemini response empty/blocked{block_reason} (Attempt {attempt + 1}).")
                 if attempt < retries - 1: time.sleep(delay); delay *= 2
                 else: print("    Warning: Empty/blocked response after max retries."); return None
        except ResourceExhausted as e:
            print(f"    Rate limit exceeded (Attempt {attempt + 1}). Retrying after {delay}s...")
            if attempt < retries - 1: time.sleep(delay); delay *= 2
            else: print("    Maximum retries reached due to rate limiting."); return None
        except Exception as e:
            print(f"    An unexpected error occurred during Gemini call: {e}"); traceback.print_exc(); return None
    return None

# --- Styling Helpers ---
def _set_table_borders(table):
    """Sets single black 0.5pt borders for all cell edges in a table."""
    try:
        border_props = {"sz": "4", "val": "single", "color": "000000", "space": "0"}
        border_keys = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        tags = {key: f"w:{key}" for key in border_keys}

        # --- Corrected tblPr access/creation ---
        tbl = table._tbl
        # Find existing tblPr or create it
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            # Insert tblPr appropriately - usually after tblGrid
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is not None:
                tblGrid.addnext(tblPr)
            else: # If no tblGrid, prepend (less common but possible)
                tbl.insert(0, tblPr)
        # --- End Correction ---

        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None: tblBorders = OxmlElement('w:tblBorders'); tblPr.append(tblBorders)
        for key in border_keys:
            border_tag = tblBorders.find(qn(tags[key]))
            if border_tag is None: border_tag = OxmlElement(tags[key]); tblBorders.append(border_tag)
            for prop, value in border_props.items(): border_tag.set(qn(f"w:{prop}"), value)
    except Exception as e: print(f"    Warning: Error applying table borders: {e}")

def _get_cell_text(cell):
    # (Standard version - no changes needed)
    if cell is None: return ""
    return "\n".join([p.text for p in cell.paragraphs])

# --- MCQ Generation Core Logic ---
def generate_mcqs(text, num_questions):
    """Generates MCQs from text using Gemini."""
    try:
        with open(config.RULES_TXT_PATH, "r", encoding='utf-8') as f: rules = f.read()
    except Exception as e: print(f"Error reading rules.txt: {e}"); return None
    prompt = f"""
Based on the rules provided in `rules.txt` (which you must follow strictly):
{rules}

**Your Task:** Generate exactly {num_questions} Multiple-Choice Questions (MCQs) from the following `#text_chunk`.

**Required Output Format (Strict):**
For EACH question, provide ALL the following components:
1.  `**Question:**` Followed by the question stem (clinical vignette or direct question), in **bold**.
2.  Five answer choices labeled `a)` to `e)`, NOT bolded, each on a new line.
3.  `**Correct Answer:**` Followed by the letter (a-e) of the correct choice, in **bold**.

**Example:**
**Question:**
**A 65-year-old male presents... Which diagnosis?**
a) Aortic stenosis
b) Mitral stenosis
c) Pulmonary embolism
d) COPD
e) VSD
**Correct Answer: b**

**Important Notes:**
*   Output ONLY the questions in the format above. NO numbering, explanations, or extra text.

**#text_chunk:**

{text}

"""
    # print("  Generating MCQs...")
    return generate_with_retry(prompt, SYSTEM_INSTRUCTION_GENERATOR)

def chunk_text(text, token_limit):
    """Splits text into chunks based on token limit."""
    sentences = re.split(r'(?<=[.!?])\s+', text); chunks = []
    current_chunk_sentences = []; current_chunk_tokens = 0
    def estimate_tokens(sentence): return len(sentence.split()) # Simple estimation
    for sentence in sentences:
        sentence_tokens = estimate_tokens(sentence)
        if current_chunk_tokens + sentence_tokens <= token_limit:
            current_chunk_sentences.append(sentence); current_chunk_tokens += sentence_tokens
        else:
            if current_chunk_sentences: chunks.append(" ".join(current_chunk_sentences))
            if sentence_tokens <= token_limit: current_chunk_sentences = [sentence]; current_chunk_tokens = sentence_tokens
            else: chunks.append(sentence); current_chunk_sentences = []; current_chunk_tokens = 0 # Handle oversized
    if current_chunk_sentences: chunks.append(" ".join(current_chunk_sentences))
    # print(f"  Text split into {len(chunks)} chunks.")
    return chunks

def calculate_num_questions(text_chunk, words_per_question):
    """Calculates how many questions per chunk."""
    word_count = len(text_chunk.split()); w_p_q = max(1, words_per_question)
    num_q = max(1, int(word_count / w_p_q))
    # print(f"  Calculated {num_q} question(s) for chunk.")
    return num_q

def verify_and_correct_mcqs(mcq_text_to_verify):
    """Verifies and corrects generated MCQs using Gemini."""
    try:
        with open(config.RULES_TXT_PATH, "r", encoding='utf-8') as f: rules = f.read()
    except Exception as e: print(f"Error reading rules.txt for verification: {e}"); return None
    prompt = f"""
Based on the rules provided in `rules.txt` (which you must follow strictly):
{rules}

**Your Task:** Review the following MCQs. Correct any violations (formatting, content, style, etc.). Ensure output perfectly matches the required format.

**Required Output Format (Strict):**
1.  `**Question:**` Stem in **bold**.
2.  Options `a)` to `e)` NOT bolded, new lines.
3.  `**Correct Answer:**` Letter (a-e) in **bold**.

**Example:**
**Question:**
**A 65-year-old male presents...**
a) Option A
b) Option B
c) Option C
d) Option D
e) Option E
**Correct Answer: b**

**Important Notes:**
*   Output ONLY the corrected questions. NO numbering, explanations, or extra text.

**MCQs to Verify and Correct:**

{mcq_text_to_verify}

"""
    # print("  Verifying and correcting MCQs...")
    return generate_with_retry(prompt, SYSTEM_INSTRUCTION_VERIFIER)

def parse_corrected_mcqs(corrected_mcq_text):
    """Parses the verified/corrected MCQ text into a list of dictionaries."""
    # Regex adjusted for robustness with potential extra newlines between options
    mcq_pattern = re.compile(
        r"^\*\*Question:\*\*\s*(.*?)\n+"       # Question stem (non-greedy), require 1+ newline
        r"^(a\).*?\n+(?:b\).*?\n+)?(?:c\).*?\n+)?(?:d\).*?\n+)?(?:e\).*?\n*))" # Options block, allow optional lines & extra newlines
        r"^\*\*Correct Answer:\s*([a-e])\*\*", # Correct answer letter
        re.DOTALL | re.MULTILINE
    )
    matches = mcq_pattern.findall(corrected_mcq_text); formatted_mcqs = []
    count = 1
    # print(f"  Parsing corrected text, found {len(matches)} potential MCQs...")
    for match in matches:
        stem, opts_block, correct_ltr = match
        clean_stem = stem.replace('**','').strip()
        # Split options more robustly, removing empty lines
        opts_lines = [opt.strip() for opt in opts_block.strip().split('\n') if opt.strip()]
        mcq_full_text = clean_stem + "\n" + "\n".join(opts_lines)
        formatted_mcqs.append({"Count": count, "MCQ": mcq_full_text, "CorrectAnswer": correct_ltr.strip()})
        count += 1
    if not formatted_mcqs: print("  Warning: Parsing found no MCQs matching expected format.")
    # else: print(f"  Successfully parsed {len(formatted_mcqs)} MCQs.")
    return formatted_mcqs

def create_mcq_markdown_table(mcq_data_list):
    """Formats a list of MCQ dicts into a Markdown table string."""
    if not mcq_data_list: return ""
    md_table = "| Question | Answer |\n|---|---|\n"
    for mcq in mcq_data_list:
        # Format question with HTML line breaks for Markdown table rendering
        q_html = mcq['MCQ'].replace('\n', '<br>')
        # Add count number before the question
        q_cell = f"{mcq['Count']}. {q_html}".replace('|', '\\|') # Escape pipes
        a_cell = mcq['CorrectAnswer'].replace('|', '\\|')
        md_table += f"| {q_cell} | {a_cell} |\n"
    return md_table

# --- DOCX Styling (MODIFIED to change row properties) ---
def apply_styling_to_mcq_docx(content_docx_path):
    """Loads DOCX, applies styles, fixes line breaks/bolding, AND disables row breaking/header repeating."""
    try:
        print("    Applying styling, fixing line breaks/bolding, and setting row properties...")
        doc = Document(content_docx_path)

        # --- Style Tables ---
        for table in doc.tables:
            _set_table_borders(table) # Apply Borders

            for row_idx, row in enumerate(table.rows):

                # --- START: Modify Row Properties ---
                try:
                    tr = row._element # Get the <w:tr> element
                    trPr = tr.get_or_add_trPr() # Get or add <w:trPr>

                    # --- 1. Disable Row Breaking Across Pages ---
                    # Find existing cantSplit element or create it
                    cantSplit = trPr.find(qn('w:cantSplit'))
                    if cantSplit is None:
                        cantSplit = OxmlElement('w:cantSplit')
                        trPr.append(cantSplit)
                    # Set attribute w:val="1" (or "true") to disable splitting
                    cantSplit.set(qn('w:val'), "1")

                    # --- 2. Disable Repeating Header Row ---
                    # Find existing tblHeader element or create it
                    tblHeader = trPr.find(qn('w:tblHeader'))
                    if tblHeader is None:
                         # Only explicitly disable if it exists or if we are the actual header row (row_idx 0)
                         # Otherwise, absence means it's disabled.
                         # Let's explicitly set it to 0 for ALL rows to be safe.
                         tblHeader = OxmlElement('w:tblHeader')
                         trPr.append(tblHeader)
                    # Set attribute w:val="0" (or "false") to disable repeating
                    tblHeader.set(qn('w:val'), "0")

                except Exception as row_prop_err:
                    # Log warning but continue processing other rows/cells
                    print(f"    Warning: Failed to set row properties for row {row_idx}: {row_prop_err}")
                # --- END: Modify Row Properties ---


                # --- Process Cells (Keep existing logic) ---
                if row_idx == 0: continue # Skip Pandoc header cell processing for now

                for cell_idx, cell in enumerate(row.cells):
                    if cell is None: continue

                    if cell_idx == 0: # Question Cell - Apply Robust Line Break Fix & Bolding
                        original_paragraphs = list(cell.paragraphs)
                        processed_texts = []
                        full_cell_text = "\n".join(p.text for p in original_paragraphs)
                        parts = re.split(r'(\s*[a-e]\))', full_cell_text)

                        if len(parts) > 1:
                            # print(f"      Reconstructing paragraphs for Row {row_idx} based on choice markers...") # Less verbose
                            stem_text = parts[0].strip()
                            if stem_text: processed_texts.append(stem_text)
                            for i in range(1, len(parts), 2):
                                marker = parts[i].strip()
                                choice_text = parts[i+1].strip() if (i+1) < len(parts) else ""
                                if marker: processed_texts.append(f"{marker} {choice_text}")
                        else:
                             # print(f"      No choice markers found for splitting in Row {row_idx}. Using original paras.") # Less verbose
                             processed_texts = [p.text for p in original_paragraphs]

                        # Clear Original Cell Content
                        if hasattr(cell, '_element'): cell._element.clear_content()
                        else: print(f"      Warning: Cannot clear cell content for Row {row_idx}.")

                        # Add new paragraphs back & style
                        for i, text in enumerate(processed_texts):
                             if not text: continue
                             new_para = cell.add_paragraph(text)
                             is_stem = (i == 0 and not re.match(r'^\s*[a-e]\)', text))
                             new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                             pPr = new_para._element.get_or_add_pPr()
                             bidi_tag = pPr.find(qn('w:bidi'));
                             if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                             bidi_tag.set(qn('w:val'), "0")
                             for run in new_para.runs:
                                run.font.name = 'Poppins'
                                r=run._element; rPr=r.get_or_add_rPr()
                                rFonts=rPr.find(qn('w:rFonts'))
                                if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.insert(0,rFonts)
                                rFonts.set(qn('w:ascii'),'Poppins'); rFonts.set(qn('w:hAnsi'),'Poppins')
                                rFonts.set(qn('w:eastAsia'),'Poppins'); rFonts.set(qn('w:cs'),'Poppins')
                                run.bold = is_stem # Bold ONLY if stem

                    elif cell_idx == 1: # Answer Cell Styling
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pPr = paragraph._element.get_or_add_pPr()
                            bidi_tag = pPr.find(qn('w:bidi'));
                            if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                            bidi_tag.set(qn('w:val'), "0")
                            for run in paragraph.runs:
                                run.font.name = 'Poppins'
                                r=run._element; rPr=r.get_or_add_rPr()
                                rFonts=rPr.find(qn('w:rFonts'))
                                if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.insert(0,rFonts)
                                rFonts.set(qn('w:ascii'),'Poppins'); rFonts.set(qn('w:hAnsi'),'Poppins')
                                rFonts.set(qn('w:eastAsia'),'Poppins'); rFonts.set(qn('w:cs'),'Poppins')
                                run.bold = False # Answer not bold

        # --- Apply Styling to Non-Table Paragraphs (Keep as before) ---
        for paragraph in doc.paragraphs:
             parent = paragraph._element.getparent(); is_in_table = False
             while parent is not None:
                 if parent.tag == qn('w:tc'): is_in_table = True; break
                 parent = parent.getparent()
             if not is_in_table:
                 for run in paragraph.runs: run.font.name = 'Poppins'
                 paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                 # ... rest of non-table paragraph styling ...
                 paragraph.paragraph_format.line_spacing = 1.0
                 p = paragraph._element; pPr = p.get_or_add_pPr()
                 bidi_tag = pPr.find(qn('w:bidi'))
                 if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                 bidi_tag.set(qn('w:val'), "0")

        # Save styled document
        doc.save(content_docx_path)
        print("    Finished styling, line break fixes, bolding, and setting row properties.")
        return True
    except Exception as e: print(f"    ERROR applying styling/fixing breaks/setting row props: {e}"); traceback.print_exc(); return False

# --- Merge, Render Template, and Final Align (Retry on Save) ---
def merge_render_align(template_path, styled_content_path, final_output_path, context):
    """Renders template, merges content, applies global LTR/Left align, retries save."""
    # (Keep this function exactly as in the previous correct answer)
    # ... (It handles rendering, merging, final alignment pass, and save retries) ...
    temp_rendered_template_path = final_output_path.replace('.docx', '_temp_rendered.docx')
    final_doc = None
    max_save_attempts = 3
    save_delay = 2 # seconds
    try:
        if not os.path.exists(template_path): print(f"  ERROR: Template not found: {template_path}."); return False
        print(f"  Rendering template with context...")
        tpl = DocxTemplate(template_path); tpl.render(context); tpl.save(temp_rendered_template_path)
        final_doc = Document(temp_rendered_template_path)
        final_doc.add_page_break()
        if not os.path.exists(styled_content_path): print(f"  ERROR: Styled content DOCX not found: {styled_content_path}."); return False
        print(f"  Appending styled content...")
        content_doc = Document(styled_content_path)
        for element in content_doc._body._element: final_doc._body._element.append(element) # Use _body._element for proper merge
        print(f"  Applying final global LTR/Left alignment...")
        for paragraph in final_doc.paragraphs:
             paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
             paragraph.paragraph_format.space_after = Pt(3)
             paragraph.paragraph_format.space_before = Pt(0)
             paragraph.paragraph_format.line_spacing = 1
             pPr = paragraph._element.get_or_add_pPr(); bidi_tag = pPr.find(qn('w:bidi'))
             if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
             bidi_tag.set(qn('w:val'), "0")
        for table in final_doc.tables:
             for row in table.rows:
                for cell in row.cells:
                    if cell is None: continue
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(3)  # Space after the paragraph (3 pt)
                        paragraph.paragraph_format.space_before = Pt(0)  # Space before the paragraph (6 pt)
                        paragraph.paragraph_format.line_spacing = 1  # Line spacing (1.5x)
                        pPr = paragraph._element.get_or_add_pPr(); bidi_tag = pPr.find(qn('w:bidi'))
                        if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                        bidi_tag.set(qn('w:val'), "0")
        for attempt in range(max_save_attempts):
            try:
                print(f"  Attempting to save final DOCX (Attempt {attempt + 1}/{max_save_attempts})...")
                final_doc.save(final_output_path)
                print(f"  Successfully merged, aligned, and saved: {os.path.basename(final_output_path)}")
                return True
            except PermissionError as pe:
                print(f"    Save failed (Permission Denied): {pe}")
                if attempt < max_save_attempts - 1: time.sleep(save_delay)
                else: print(f"  ERROR: Failed to save DOCX after {max_save_attempts} attempts (PermissionError)."); return False
            except Exception as save_err: print(f"  Error during final save: {save_err}"); traceback.print_exc(); return False
        return False
    except Exception as e: print(f"  Error during final merge/render/align: {e}"); traceback.print_exc(); return False
    finally:
        if os.path.exists(temp_rendered_template_path):
            try: os.remove(temp_rendered_template_path)
            except OSError as e: print(f"  Warning: Could not remove temp file {temp_rendered_template_path}: {e}")


# --- Main Orchestration Function ---
def create_mcqs_and_process(md_path, output_csv_path, output_docx_path, template_path):
    """Reads MD, generates/verifies MCQs, saves CSV, creates/styles/merges DOCX."""
    # (Keep this function exactly as in the previous correct answer)
    # ... (It calls the generation, styling, and merging steps in order) ...
    print(f"Processing MCQs for: {os.path.basename(md_path)}")
    base_name = os.path.splitext(os.path.basename(md_path))[0]
    if base_name.endswith("_extracted"): base_name = base_name[:-10]
    lecture_name = base_name.replace('_', ' ').replace('-', ' ').title().upper()
    context = {'lecture_name': lecture_name}
    print(f"  Using Lecture Name: {lecture_name}")
    temp_md_path = output_csv_path.replace('.csv', '_temp.md')
    temp_styled_content_docx_path = output_csv_path.replace('.csv', '_temp_styled_content.docx')
    try:
        with open(md_path, 'r', encoding='utf-8') as f: text = f.read()
        if not text.strip(): print("  Warning: MD empty."); return False
        chunks = chunk_text(text, config.TOKEN_LIMIT)
        if not chunks: print("  Warning: No text chunks."); return False
        all_raw_mcqs = []
        for i, chunk in enumerate(chunks):
            num_q = calculate_num_questions(chunk, config.WORDS_PER_QUESTION)
            raw_mcq = generate_mcqs(chunk, num_q)
            if raw_mcq: all_raw_mcqs.append(raw_mcq)
        if not all_raw_mcqs: print("  No MCQs generated."); return False
        combined_raw = "\n\n".join(all_raw_mcqs)
        corrected_mcq = verify_and_correct_mcqs(combined_raw)
        if not corrected_mcq: print("  MCQ verification failed."); return False
        parsed_mcqs = parse_corrected_mcqs(corrected_mcq)
        if not parsed_mcqs: print("  Failed to parse verified MCQs."); return False
        try:
            df = pd.DataFrame(parsed_mcqs); df.to_csv(output_csv_path, index=False, encoding='utf-8')
            print(f"  Saved {len(parsed_mcqs)} MCQs to CSV: {os.path.basename(output_csv_path)}")
        except Exception as e: print(f"  Error saving CSV: {e}"); traceback.print_exc(); return False
        mcq_md_table = create_mcq_markdown_table(parsed_mcqs)
        if not mcq_md_table: print("  Failed to create MD table string."); return False
        with open(temp_md_path, 'w', encoding='utf-8') as f_md: f_md.write(mcq_md_table)
        print(f"  Converting MD table to DOCX via Pandoc...")
        pypandoc_args = ['--wrap=none']
        pandoc_path = getattr(config, 'PANDOC_PATH', None)
        if pandoc_path and os.path.exists(pandoc_path):
             try: pypandoc.pandoc_download.check_pandoc_path(pandoc_path); pypandoc_args.extend(['--pandoc-binary', pandoc_path])
             except Exception as pe: print(f"  Warning: Issue with custom Pandoc path {pandoc_path}: {pe}")
        pypandoc.convert_file(temp_md_path, 'docx', outputfile=temp_styled_content_docx_path, extra_args=pypandoc_args)
        styling_success = apply_styling_to_mcq_docx(temp_styled_content_docx_path)
        if not styling_success: print("  ERROR: Styling/Line Break Fix failed."); return False
        merge_success = merge_render_align(
            template_path=template_path,
            styled_content_path=temp_styled_content_docx_path,
            final_output_path=output_docx_path,
            context=context
        )
        return merge_success
    except pypandoc.PandocMissingError: print("ERROR: Pandoc not found."); traceback.print_exc(); return False
    except FileNotFoundError as e: print(f"Error: Input file not found during process: {e}"); traceback.print_exc(); return False
    except Exception as e: print(f"Unexpected error in create_mcqs_and_process: {e}"); traceback.print_exc(); return False
    finally:
        print("Cleaning up temporary files...")
        for temp_file in [temp_md_path, temp_styled_content_docx_path]:
             if os.path.exists(temp_file):
                 try: os.remove(temp_file)
                 except OSError as e: print(f"  Warning: Could not remove temp file {temp_file}: {e}")