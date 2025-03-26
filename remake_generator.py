# remake_generator.py

import google.generativeai as genai
import config # Ensure config is imported
import os
import time
import json # <-- Added
import re   # <-- Added
from google.api_core.exceptions import ResourceExhausted
from docxtpl import DocxTemplate
import pypandoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import lxml.etree as etree
from docx.oxml.shared import OxmlElement, qn
from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
import traceback # Import traceback for detailed error logging

# --- Helper Functions qn, _set_table_borders (Keep as before) ---
def qn(tag_name):
    """Stands for 'qualified name', a utility function for lxml."""
    import docx.oxml.ns
    prefix, tag = tag_name.split(':')
    uri = '{%s}%s' % (docx.oxml.ns.nsmap[prefix], tag)
    return uri

# --- Helper Function to set table borders ---
def _set_table_borders(table):
    """Sets single black 0.5pt borders for all cell edges in a table."""
    try:
        border_props = {
            "sz": "4", "val": "single", "color": "000000", "space": "0"
        }
        border_keys = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        tags = {key: f"w:{key}" for key in border_keys}

        tbl = table._tbl # Get the underlying CT_Tbl object

        # --- Correctly find or create tblPr ---
        tblPr = tbl.find(qn('w:tblPr')) # Find existing tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr') # Create if not found
            # Insert tblPr as the first child of tbl (common practice)
            tbl.insert(0, tblPr)
        # --- End of correction for tblPr ---

        # --- Correctly find or create tblBorders ---
        tblBorders = tblPr.find(qn('w:tblBorders')) # Find existing tblBorders within tblPr
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders') # Create if not found
            tblPr.append(tblBorders) # Append to tblPr
        # --- End of correction for tblBorders ---

        # Apply individual border settings (this part was okay)
        for key in border_keys:
            border_tag = tblBorders.find(qn(tags[key]))
            if border_tag is None:
                border_tag = OxmlElement(tags[key])
                tblBorders.append(border_tag)
            for prop, value in border_props.items():
                 border_tag.set(qn(f"w:{prop}"), value)

    except Exception as e:
        print(f"Error applying borders to table: {e}")
        raise # Re-raise the exception after printing

# --- Helper function _get_cell_text (Keep as before) ---
# def _get_cell_text(cell): ...

# --- Gemini API Call with Retry (MODIFIED FOR JSON) ---
def generate_with_retry(prompt, retries=5, delay=5, expect_json=False): # Added expect_json flag
    """Retries the generation request, handling JSON expectation."""
    genai.configure(api_key=config.API_KEY)

    gen_config = config.generation_config.copy()
    if expect_json:
        gen_config["response_mime_type"] = "application/json"
    else: # Ensure it's text/plain if not JSON
        gen_config["response_mime_type"] = "text/plain"

    model = genai.GenerativeModel(
        model_name=config.GEMINI_MODEL,
        generation_config=gen_config,
        safety_settings=config.safety_settings
    )

    for attempt in range(retries):
        try:
            print(f"remake Gen: Sending request to Gemini (attempt {attempt + 1}, expecting {'JSON' if expect_json else 'Text'})...")
            response = model.generate_content(prompt)
            print("remake Gen: Received response from Gemini.")
            # Simple check: if JSON was expected but we got blocked, return None early
            if expect_json:
                try:
                     # Check for blocking before accessing text
                     if response.prompt_feedback.block_reason:
                          print(f"remake Gen: Content blocked due to: {response.prompt_feedback.block_reason}")
                          return None
                except Exception: pass # Ignore if feedback check fails
            return response.text # Return the text part (should be JSON string if expect_json)
        except ResourceExhausted as e:
            print(f"remake Gen: Rate limit exceeded, retrying in {delay}s... ({attempt + 1}/{retries})")
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            print(f"remake Gen: An unexpected error occurred during Gemini call: {e}")
            # ... (Error feedback logging) ...
            try: # Safely check feedback
                 if response and hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
                      print(f"remake Gen: Content blocked due to: {response.prompt_feedback.block_reason}")
                 elif response and hasattr(response, 'candidates') and response.candidates and response.candidates[0].finish_reason != 'STOP':
                     print(f"remake Gen: Generation stopped potentially due to: {response.candidates[0].finish_reason}")
            except Exception as feedback_err:
                 print(f"remake Gen: Error accessing feedback details: {feedback_err}")
            # Retry logic
            if attempt < retries - 1:
                print(f"remake Gen: Retrying after error in {delay}s...")
                time.sleep(delay)
                delay *= 2
            else:
                print("remake Gen: Maximum retries reached or fatal error.")
                traceback.print_exc()
                if expect_json: return None # Return None on failure when expecting JSON
                raise e # Re-raise if not expecting JSON or want failure propagation

    print("remake Gen: Max retries reached after ResourceExhausted errors.")
    return None # Return None if loop finishes without success

# --- UPDATED: Generate remake JSON Content using Gemini ---
def generate_remake_json_content(text):
    """Generates the initial structured 'remake' as JSON using Gemini."""
    print("Remake Gen: Generating initial structured JSON via Gemini...")
    prompt = f"""
**Role:** You are an expert AI medical content structuring specialist. Your task is to meticulously analyze medical lecture text and **reconstruct** it into a detailed, structured JSON format, preserving **high fidelity** to the original content.

**Audience:** Medical students needing comprehensive, organized notes.

**Task:** Analyze the provided medical lecture text. Break it down into logical sections. For each section, further break down the information into key concepts or subtopics and their associated details. Generate a JSON output representing this detailed reconstruction. **The goal is restructuring for clarity, NOT summarization.**

**JSON Output Structure:**
The output MUST be a single JSON list `[...]` containing section objects. Each section object MUST have the following structure:
```json
{{
  "title": "Section Title (e.g., Pathophysiology of Disease X)", // String: Clear, descriptive heading for the section
  "content": [                                                // List of key-detail pairs for this section
    {{
      "key_point": "Subtopic or Concept 1",                  // String: The main concept/term/subheading within the section
      "details": "Comprehensive details from the text related to Key Point 1. Include definitions, explanations, examples, mechanisms, specific data, etc. Preserve original meaning and detail." // String: All relevant details for this key point
    }},
    {{
      "key_point": "Subtopic or Concept 2",
      "details": "Comprehensive details from the text related to Key Point 2..."
    }}
    // ... more key-detail pairs as needed for the section
  ]
}}
Use code with caution.
Python
Each section MUST use the "content" list containing {{ "key_point": ..., "details": ... }} objects.

The "details" string should contain all relevant information from the original text pertaining to that "key_point". Embed lists or steps within the details string if necessary (e.g., using bullet points or numbered lists within the text).

Content Rules:

High Fidelity: Capture all essential information, including details, definitions, examples, classifications, mechanisms, specific values, etc., presented in the source text. Avoid simplification or omission unless the information is truly trivial or redundant within the same context.

Structure: Organize information logically under appropriate section titles. Within each section, group related information under distinct key_point entries.

Accuracy: MUST accurately reflect the meaning and terminology of the source text.

No External Information: Do NOT add information not present in the original text.

Conciseness (Titles only): Section titles and key_point strings should be clear and reasonably concise headings, but the "details" string MUST be comprehensive.

Example JSON Output:

[
  {{
    "title": "Pharmacokinetics: Absorption",
    "content": [
      {{
        "key_point": "Definition",
        "details": "Process by which a drug enters the systemic circulation from the site of administration."
      }},
      {{
        "key_point": "Factors Affecting Absorption",
        "details": "Includes: Route of administration (e.g., oral, IV, topical), drug properties (lipid solubility, ionization), blood flow to absorption site, surface area available, contact time, presence of food or other drugs, first-pass metabolism (especially for oral drugs)."
      }},
      {{
        "key_point": "Bioavailability (F)",
        "details": "Fraction of administered drug that reaches systemic circulation unchanged. IV route has F=100%. Oral bioavailability is often lower due to incomplete absorption and first-pass effect. Calculated as (AUC oral / AUC IV) * (Dose IV / Dose oral)."
      }}
    ]
  }},
  {{
    "title": "Pharmacokinetics: Distribution",
    "content": [
      {{
        "key_point": "Definition",
        "details": "Reversible transfer of a drug from the bloodstream to tissues."
      }},
      {{
        "key_point": "Volume of Distribution (Vd)",
        "details": "Apparent volume into which a drug distributes. Relates the amount of drug in the body to the concentration in plasma (Vd = Amount of drug in body / Plasma drug concentration). High Vd suggests extensive tissue distribution (e.g., lipid-soluble drugs). Low Vd suggests confinement to plasma (e.g., highly protein-bound drugs)."
      }},
      {{
          "key_point": "Factors Affecting Distribution",
          "details": "Blood flow to tissues, drug's ability to exit vasculature (lipid solubility, protein binding), drug's ability to enter cells."
      }}
    ]
  }}
]
Use code with caution.
Json
Input Text:
{text}
Final Output Instruction:
Generate ONLY the JSON list [...] based on the requirements above. Do not include any other text, comments, or markdown formatting like ```json.
"""
    json_string = generate_with_retry(prompt, expect_json=True) # Expect JSON response
    if not json_string:
        print("remake Gen: ERROR - Failed to get valid JSON response from Gemini for initial generation.")
        return None

    try:
        print("remake Gen: Parsing Gemini JSON response...")
        # Clean potential markdown formatting
        if json_string.strip().startswith("```json"):
             json_string = json_string.strip()[7:-3].strip()
        elif json_string.strip().startswith("```"):
             json_string = json_string.strip()[3:-3].strip()
        remake_data = json.loads(json_string)
        # Basic validation: Check if it's a list
        if not isinstance(remake_data, list):
            print("remake Gen: ERROR - Parsed JSON is not a list as expected.")
            print("--- Received Data ---")
            print(remake_data)
            print("---------------------")
            return None
        print("remake Gen: Initial JSON structure parsed successfully.")
        return remake_data # Return the parsed list
    except json.JSONDecodeError as e:
        print(f"remake Gen: ERROR - Failed to decode JSON response: {e}")
        print("--- Gemini Response Text (raw) ---")
        print(json_string)
        print("----------------------------------")
        return None
    except Exception as e:
        print(f"remake Gen: An unexpected error occurred parsing JSON: {e}")
        traceback.print_exc()
        return None


# --- UPDATED: Verify and Correct remake JSON ---
def verify_and_correct_remake_json(original_text, remake_json_to_verify):
    """
    Compares the generated remake JSON to the original_text for fidelity,
    completeness, and adherence to JSON structure, refining it if necessary.
    Returns the *verified* JSON data (list).
    """
    print("Remake Gen: Starting JSON remake verification/correction...")
    try:
        json_string_to_verify = json.dumps(remake_json_to_verify, indent=2)
    except Exception as e:
        print(f"Remake Gen: ERROR - Could not serialize JSON for verification prompt: {e}")
        return None

    verification_rules = """
**Core Remake Verification Goals:**
1.  **Fidelity & Completeness:** The JSON remake MUST accurately and **comprehensively** represent *all essential information* present in the `ORIGINAL_TEXT`. Check for omissions of important details, definitions, examples, mechanisms, etc. Add missing information to the appropriate `"details"` field. Correct any factual inaccuracies.
2.  **No External Information:** Ensure NO information exists in the JSON that cannot be directly supported by the `ORIGINAL_TEXT`. Remove any external additions.
3.  **Structure Adherence:** The output MUST be a valid JSON list `[...]`. Each object in the list MUST have `"title"` (string) and `"content"` (list). The `"content"` list MUST contain only objects with `"key_point"` (string) and `"details"` (string). Correct any structural deviations *strictly* according to this format.
4.  **Logical Grouping:** Ensure `key_point` entries logically group related `details` within each section `title`. Minor reorganization for better logical flow is acceptable if fidelity is maintained.
5.  **Conciseness (Titles):** `title` and `key_point` should be clear headings, but `"details"` must remain comprehensive.
"""

    prompt = f"""
**Role:** You are an expert AI medical content editor focused on verifying the **fidelity** and structure of reconstructed content.

**Objective:** Analyze the `JSON_REMAKE_TO_VERIFY`. Compare it meticulously against the `ORIGINAL_TEXT` based on the `VERIFICATION_GOALS`. Ensure the JSON is an accurate, complete, and structurally correct representation of the original material.

{verification_rules}

---
**Input 1: ORIGINAL_TEXT** (The full source material)

{original_text}

---
**Input 2: JSON_REMAKE_TO_VERIFY** (The JSON generated previously)
```json
{json_string_to_verify}
Use code with caution.
Verification & Refinement Task:
Review the JSON_REMAKE_TO_VERIFY against the ORIGINAL_TEXT.

Check Fidelity/Completeness: Add any missing essential details from the ORIGINAL_TEXT into the relevant "details" fields. Correct any misrepresentations.

Check External Info: Remove information not derivable from ORIGINAL_TEXT.

Check JSON Structure: Ensure the entire output is a valid JSON list following the exact specified structure (list -> objects with "title" and "content" -> "content" is a list of objects with "key_point" and "details"). Fix any structural errors rigorously.

Check Logical Grouping: Ensure key_point/details pairs are logical.

Output Instruction:
Output ONLY the final, verified, and potentially refined JSON list [...].

If refinements were made, output the improved JSON list.

If the original JSON_REMAKE_TO_VERIFY met all goals perfectly, output it exactly as provided.

Do NOT include any explanations, comments, confirmations, or conversational text. Your entire response must be the final JSON list content, starting with [ and ending with ].
"""

    verified_json_string = generate_with_retry(prompt, expect_json=True) # Expect JSON back
    if not verified_json_string:
        print("remake Gen: ERROR - Failed to get valid JSON response from Gemini for verification.")
        return None

    try:
        print("remake Gen: Parsing verified Gemini JSON response...")
        # Clean potential markdown formatting
        if verified_json_string.strip().startswith("```json"):
             verified_json_string = verified_json_string.strip()[7:-3].strip()
        elif verified_json_string.strip().startswith("```"):
             verified_json_string = verified_json_string.strip()[3:-3].strip()
        verified_remake_data = json.loads(verified_json_string)
        # Basic validation
        if not isinstance(verified_remake_data, list):
            print("remake Gen: ERROR - Verified JSON is not a list as expected.")
            print("--- Received Data ---")
            print(verified_remake_data)
            print("---------------------")
            return None
        print("remake Gen: Verified JSON structure parsed successfully.")
        return verified_remake_data # Return the parsed verified list
    except json.JSONDecodeError as e:
        print(f"remake Gen: ERROR - Failed to decode verified JSON response: {e}")
        print("--- Verified Gemini Response Text (raw) ---")
        print(verified_json_string)
        print("-----------------------------------------")
        return None
    except Exception as e:
        print(f"remake Gen: An unexpected error occurred parsing verified JSON: {e}")
        traceback.print_exc()
        return None

# --- NEW: Function to Convert JSON remake to Markdown ---
def json_to_markdown_remake(remake_data):
    """Converts the structured remake data (list of sections) into a Markdown string."""
    if not remake_data:
        print("remake Gen: Warning - JSON remake data is empty, generating empty Markdown.")
        return ""

    md_parts = []
    for section in remake_data:
        title = section.get("title", "Untitled Section")
        type = section.get("type")
        content = section.get("content")

        # Add Heading
        md_parts.append(f"## {title}\n") # Use H2 for sections

        if type == "paragraph" and isinstance(content, str):
            md_parts.append(f"{content}\n")
        elif type == "list" and isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    md_parts.append(f"- {item}\n")
            md_parts.append("\n") # Add space after list
        elif type == "table" and isinstance(content, list) and content:
            # Check if content is a list of dicts with expected keys
            if all(isinstance(row, dict) and "key_point" in row and "details" in row for row in content):
                # Create Markdown Table
                md_parts.append("| Key Point | Details |\n") # Fixed headers
                md_parts.append("|---|---|\n")
                for row in content:
                    # Escape pipe characters within cell content
                    key_point_cell = str(row.get("key_point", "")).replace('|', '\\|')
                    details_cell = str(row.get("details", "")).replace('|', '\\|')
                    md_parts.append(f"| {key_point_cell} | {details_cell} |\n")
                md_parts.append("\n") # Add space after table
            else:
                print(f"remake Gen: Warning - Section '{title}' has type 'table' but content format is invalid. Skipping table generation.")
                md_parts.append(f"[Content for section '{title}' intended as table, but format was invalid]\n")
        else:
            print(f"remake Gen: Warning - Section '{title}' has unknown type '{type}' or invalid content. Skipping content.")
            md_parts.append(f"[Content for section '{title}' has unknown type '{type}' or invalid content]\n")

        md_parts.append("\n") # Add space between sections

    return "".join(md_parts)


# --- UPDATED: Function to Apply Styling (Removed Row Merging) ---
def apply_styling_to_remake_docx(docx_path):
    """Loads a DOCX file, applies table borders, paragraph styling, and saves in place."""
    print(f"remake Styling: Applying styling to: {docx_path}")
    try:
        doc = Document(docx_path)

        # --- Style Tables (Borders and Font) ---
        print(f"  Found {len(doc.tables)} tables.")
        for table_idx, table in enumerate(doc.tables):
            try:
                # --- Apply Borders ---
                _set_table_borders(table)

                # --- Apply Font & Bold First Column ---
                # Pandoc MD tables usually have a header row, skip it for bolding data
                has_header = (len(table.rows) > 1 and
                                table.rows[0].cells and len(table.rows[0].cells) > 0 and
                                table.rows[0].cells[0].text.strip().lower() == "key point") # Check header based on our generation

                for row_idx, row in enumerate(table.rows):
                    # Bold first column data cells (skip header)
                    if has_header and row_idx == 0:
                         pass # Skip actual header row if detected
                    elif len(row.cells) > 0:
                         for paragraph in row.cells[0].paragraphs:
                               for run in paragraph.runs:
                                   run.bold = True

                    # Apply font to all cells
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Poppins'
                                # Apply Oxml font settings (optional but good practice)
                                r = run._element; rPr = r.get_or_add_rPr()
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
                                rFonts.set(qn('w:ascii'), 'Poppins'); rFonts.set(qn('w:hAnsi'), 'Poppins')
                                rFonts.set(qn('w:eastAsia'), 'Poppins'); rFonts.set(qn('w:cs'), 'Poppins')

            except Exception as table_style_error:
                print(f"  ERROR: Failed styling table {table_idx+1}: {table_style_error}")
                traceback.print_exc()

        # --- Apply Other Formatting (non-table paragraphs) ---
        for paragraph in doc.paragraphs:
             # Font (Apply Poppins)
             for run in paragraph.runs:
                run.font.name = 'Poppins'
                r = run._element; rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Poppins'); rFonts.set(qn('w:hAnsi'), 'Poppins')
                rFonts.set(qn('w:eastAsia'), 'Poppins'); rFonts.set(qn('w:cs'), 'Poppins')

             # Alignment, spacing, list indents (Keep existing logic)
             paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
             paragraph.paragraph_format.line_spacing = 1.0
             if paragraph.style.name.startswith('Heading'):
                 paragraph.paragraph_format.left_indent = Inches(0)
                 paragraph.paragraph_format.space_before = Pt(12)
                 paragraph.paragraph_format.space_after = Pt(6)
             elif paragraph.style.name.startswith('List'):
                  # Let Pandoc handle list indentation, just ensure alignment/direction
                  pass # Or apply minimal indent if needed: paragraph.paragraph_format.left_indent = Inches(0.25)
             else: # Regular paragraphs
                  paragraph.paragraph_format.left_indent = Inches(0) # Often 0 indent is fine for summaries
                  paragraph.paragraph_format.first_line_indent = Inches(0)

             # LTR direction (Keep existing logic)
             p = paragraph._element; pPr = p.get_or_add_pPr()
             bidi_tag = pPr.find(qn('w:bidi'))
             if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
             bidi_tag.set(qn('w:val'), "0")

        # Save the styled document back to the same path
        doc.save(docx_path)
        print(f"  Styling applied successfully to {docx_path}")
        return True

    except Exception as e:
        print(f"ERROR applying styling to {docx_path}: {e}")
        traceback.print_exc()
        return False


# --- Merge Template and Content (Keep as before) ---
def merge_template_and_styled_content(intro_docx_path, styled_content_path, final_output_path):
    # ... (No changes needed here) ...
    print(f"Merging '{intro_docx_path}' and '{styled_content_path}' into '{final_output_path}'")
    try:
        intro_doc = Document(intro_docx_path)
        intro_doc.add_page_break()
        content_doc = Document(styled_content_path)
        for element in content_doc.element.body:
            intro_doc.element.body.append(element)

        # Final LTR/Alignment Pass
        print("  Applying final LTR/Alignment pass to merged document...")
        for paragraph in intro_doc.paragraphs:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = paragraph._element; pPr = p.get_or_add_pPr()
            bidi_tag = pPr.find(qn('w:bidi'))
            if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
            bidi_tag.set(qn('w:val'), "0")
        for table in intro_doc.tables:
             for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p = paragraph._element; pPr = p.get_or_add_pPr()
                        bidi_tag = pPr.find(qn('w:bidi'))
                        if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                        bidi_tag.set(qn('w:val'), "0")

        intro_doc.save(final_output_path)
        print(f"  Successfully merged and saved to {final_output_path}")
        return True
    except Exception as e:
        print(f"ERROR during merging/final saving: {e}")
        traceback.print_exc()
        return False


# remake_generator.py

# ... (other imports, helper functions, generate_with_retry, JSON generation/verification) ...

# --- NEW: Function to Convert JSON Remake to Markdown ---
def json_to_markdown_remake(remake_data):
    """Converts the structured remake data (list of sections) into Markdown,
       assuming each section's content is for a table."""
    if not remake_data:
        print("Remake Gen: Warning - JSON remake data is empty, generating empty Markdown.")
        return ""

    md_parts = []
    for section in remake_data:
        title = section.get("title", "Untitled Section")
        content_list = section.get("content") # Get the list of key-detail pairs

        # Add Heading (Use H2 for sections)
        md_parts.append(f"## {title}\n")

        # --- Assume content is always for a table based on remake prompt structure ---
        if isinstance(content_list, list) and content_list:
            # Check if content is a list of dicts with expected keys
            if all(isinstance(row, dict) and "key_point" in row and "details" in row for row in content_list):
                # Create Markdown Table
                md_parts.append("| Key Point | Details |\n") # Standard headers for remake
                md_parts.append("|---|---|\n")
                for row in content_list:
                    # Escape pipe characters within cell content
                    key_point_cell = str(row.get("key_point", "")).replace('|', '\\|')
                    details_cell = str(row.get("details", "")).replace('|', '\\|')

                    # IMPORTANT: Replace newlines in 'details' with <br> for Pandoc
                    # This helps preserve line breaks when converting MD table to DOCX
                    details_cell_md = details_cell.replace('\n', '<br>')

                    md_parts.append(f"| {key_point_cell} | {details_cell_md} |\n")
                md_parts.append("\n") # Add space after table
            else:
                # Content list is present but items don't have the right keys
                print(f"Remake Gen: Warning - Section '{title}' content list items have invalid format. Skipping table generation.")
                md_parts.append(f"[Content for section '{title}' format invalid (expected list of {{'key_point': ..., 'details': ...}})]\n\n")
        elif content_list is None:
            print(f"Remake Gen: Warning - Section '{title}' is missing 'content' list. Skipping.")
            md_parts.append(f"[Section '{title}' is missing content]\n\n")
        else: # Content is not a non-empty list (e.g., empty list [])
            print(f"Remake Gen: Warning - Section '{title}' has empty or invalid 'content'. Skipping.")
            md_parts.append(f"[Section '{title}' has empty or invalid content]\n\n")
        # --- End of Table Logic ---

        # md_parts.append("\n") # Add extra space between sections if desired (already added one after table/message)

    return "".join(md_parts)

# --- Adapt create_remake function ---
def create_remake(input_md_path, output_docx_path, template_path):
    """Creates a remake DOCX from Markdown via JSON intermediate."""
    print(f"Starting remake process for: {input_md_path}")
    try:
        # Read Original Text
        print("Reading original Markdown file...")
        with open(input_md_path, 'r', encoding='utf-8') as f:
            original_md_text = f.read()
        if not original_md_text.strip():
            print("Warning: Input Markdown file is empty.")
            return False

        # Prepare base name
        base_name = os.path.splitext(os.path.basename(input_md_path))[0]
        if base_name.endswith("_extracted"): base_name = base_name[:-10]
        print(f"Using base name: {base_name}")

        # 1. Generate Initial JSON Remake Content
        initial_remake_json = generate_remake_json_content(original_md_text) # Use remake-specific function
        if initial_remake_json is None:
            print("ERROR: Failed to generate initial remake JSON from Gemini.")
            return False
        print("Initial Gemini JSON generation successful.")

        # 2. Verify and Correct the Initial JSON Remake
        print("Verifying and correcting generated JSON remake...")
        verified_remake_json = verify_and_correct_remake_json(original_md_text, initial_remake_json) # Use remake-specific function
        if verified_remake_json is None:
            print("ERROR: Failed during remake JSON verification/correction step.")
            return False
        print("Verification/correction of JSON step complete.")

        # 3. Convert Verified JSON to Markdown String *** USING THE CORRECT FUNCTION ***
        print("Converting verified JSON to Markdown string...")
        final_markdown_content = json_to_markdown_remake(verified_remake_json) # Use the new remake function
        if not final_markdown_content:
             print("Warning: Generated Markdown content is empty.")
             # return False # Decide if empty output is failure
        print("Markdown string generated successfully.")

        # 4. Process Markdown to Styled and Merged DOCX
        print("Starting DOCX processing (Pandoc, Styling, Merging)...")
        # Use a suitable styling function - likely similar to summary but maybe named differently
        # Ensure process_remake_md_to_docx calls a styling function that handles these tables correctly
        success = process_remake_md_to_docx( # Assuming you have/will create this orchestrator
            md_content=final_markdown_content,
            lecture_name=base_name,
            template_path=template_path,
            output_docx_path=output_docx_path
        )

        if success:
            print(f"--- remake process completed successfully for: {input_md_path} ---")
        else:
            print(f"--- remake process failed for: {input_md_path} ---")

        return success

    # ... (rest of exception handling like in create_summary) ...
    except FileNotFoundError:
        print(f"ERROR: Input Markdown file '{input_md_path}' or template file '{template_path}' not found.")
        traceback.print_exc()
        return False
    except Exception as e:
        print(f"An UNEXPECTED error occurred during the remake generation process: {e}")
        traceback.print_exc()
        return False


def apply_styling_to_remake_docx(docx_path):
     """Loads DOCX, applies table borders, paragraph styling for Remake."""
     # This function will be VERY similar to apply_styling_to_summary_docx
     # but without the row merging and ensuring it handles the Key Point/Details headers
     print(f"Remake Styling: Applying styling to: {docx_path}")
     try:
        doc = Document(docx_path)
        # --- Style Tables (Borders and Font) ---
        print(f"  Found {len(doc.tables)} tables.")
        for table_idx, table in enumerate(doc.tables):
            try:
                _set_table_borders(table) # Apply Borders

                # Apply Font & Bold First Column (Key Point)
                has_header = (len(table.rows) > 1 and
                                table.rows[0].cells and len(table.rows[0].cells) > 0 and
                                table.rows[0].cells[0].text.strip().lower() == "key point")

                for row_idx, row in enumerate(table.rows):
                    # Bold first column data cells
                    if has_header and row_idx == 0: pass # Skip header
                    elif len(row.cells) > 0:
                         for paragraph in row.cells[0].paragraphs:
                             for run in paragraph.runs: run.bold = True
                    # Apply font to all cells
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # --- Fix <br> back to newlines ---
                            # Pandoc might convert <br> to multiple runs or paragraphs.
                            # This basic replacement handles simple cases. May need refinement.
                            if '<br>' in paragraph.text:
                                text_parts = paragraph.text.split('<br>')
                                paragraph.text = text_parts[0] # Keep first part in current para
                                for part in text_parts[1:]:
                                     # Insert new paragraph *before* current one for remaining parts? No, add *after*.
                                     # Or, better: replace <br> with newline char within the paragraph?
                                     # Simplest approach for now: Replace in text and let styling apply.
                                     # This might not create actual line breaks IN the cell, Pandoc/Word handles it.
                                     # Let's try replacing with \n in the run text
                                     pass # Need a more robust way to handle <br> -> newline in docx

                            # Apply Poppins font
                            for run in paragraph.runs:
                                run.font.name = 'Poppins'
                                # Apply Oxml font settings
                                r = run._element; rPr = r.get_or_add_rPr()
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
                                rFonts.set(qn('w:ascii'), 'Poppins'); rFonts.set(qn('w:hAnsi'), 'Poppins')
                                rFonts.set(qn('w:eastAsia'), 'Poppins'); rFonts.set(qn('w:cs'), 'Poppins')

            except Exception as table_style_error:
                print(f"  ERROR: Failed styling table {table_idx+1}: {table_style_error}")
                traceback.print_exc()
        # --- Apply Other Formatting (non-table paragraphs) ---
        # ... (Paragraph styling logic - likely identical to summary styling) ...
        for paragraph in doc.paragraphs:
            # Font
            for run in paragraph.runs:
                run.font.name = 'Poppins'
                r=run._element; rPr=r.get_or_add_rPr()
                rFonts=rPr.find(qn('w:rFonts'))
                if rFonts is None: rFonts=OxmlElement('w:rFonts'); rPr.insert(0,rFonts)
                rFonts.set(qn('w:ascii'),'Poppins'); rFonts.set(qn('w:hAnsi'),'Poppins')
                rFonts.set(qn('w:eastAsia'),'Poppins'); rFonts.set(qn('w:cs'),'Poppins')
            # Alignment, spacing, etc.
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            if paragraph.style.name.startswith('Heading'):
                paragraph.paragraph_format.left_indent = Inches(0)
                # ... spacing ...
            else: # Regular/List paragraphs
                paragraph.paragraph_format.left_indent = Inches(0)
            # LTR
            p=paragraph._element; pPr=p.get_or_add_pPr()
            bidi=pPr.find(qn('w:bidi'))
            if bidi is None: bidi=etree.SubElement(pPr, qn('w:bidi'))
            bidi.set(qn('w:val'), "0")

        doc.save(docx_path)
        print(f"  Styling applied successfully to {docx_path}")
        return True
     except Exception as e:
        print(f"ERROR applying styling to {docx_path}: {e}")
        traceback.print_exc()
        return False


def process_remake_md_to_docx(md_content, lecture_name, template_path, output_docx_path):
    """Orchestrates MD save, Pandoc conversion, styling, and merging for Remake."""
    # This is almost identical to process_summary_md_to_docx, just calls apply_styling_to_remake_docx
    temp_md_path = f"temp_remake_{lecture_name}.md"
    temp_intro_path = f"temp_remake_intro_{lecture_name}.docx"
    temp_content_unstyled_path = f"temp_remake_content_unstyled_{lecture_name}.docx"
    try:
        # 1. Render template intro
        print("Loading and rendering template...")
        doc_template = DocxTemplate(template_path)
        context = {'lecture_name': lecture_name}
        doc_template.render(context)
        doc_template.save(temp_intro_path)
        # 2. Save generated Markdown content
        print(f"Saving generated Markdown to '{temp_md_path}'...")
        with open(temp_md_path, 'w', encoding='utf-8') as f_md: f_md.write(md_content)
        # 3. Convert MD to DOCX via Pandoc
        print(f"Converting MD to DOCX via Pandoc: '{temp_md_path}' -> '{temp_content_unstyled_path}'...")
        pypandoc_args = ['--wrap=none']
        # Add pandoc path logic if needed
        pandoc_path = getattr(config, 'PANDOC_PATH', None)
        if pandoc_path and os.path.exists(pandoc_path):
             try: pypandoc_args.extend(['--pandoc-binary', pandoc_path]); print(f"  Using custom Pandoc: {pandoc_path}")
             except Exception as pe: print(f"  Warning: Issue using custom Pandoc path {pandoc_path}: {pe}.")
        pypandoc.convert_file(temp_md_path, 'docx', outputfile=temp_content_unstyled_path, extra_args=pypandoc_args)
        print("  Pandoc conversion successful.")
        # 4. Apply styling *** USING REMAKE STYLING FUNCTION ***
        styling_success = apply_styling_to_remake_docx(temp_content_unstyled_path) # Call the remake styler
        if not styling_success: print("ERROR: Failed to apply styling."); return False
        # 5. Merge
        merge_success = merge_template_and_styled_content(temp_intro_path, temp_content_unstyled_path, output_docx_path) # Merge function is generic
        if not merge_success: print("ERROR: Failed to merge template and styled content."); return False
        print(f"Remake DOCX generated successfully: {output_docx_path}")
        return True
    # ... (exception handling identical to process_summary_md_to_docx) ...
    except pypandoc.PandocMissingError: print("ERROR: Pandoc executable not found."); traceback.print_exc(); return False
    except FileNotFoundError as fnf_error: print(f"ERROR: File not found: {fnf_error}"); traceback.print_exc(); return False
    except Exception as e: print(f"An CRITICAL error occurred during DOCX orchestration: {e}"); traceback.print_exc(); return False
    finally: # Cleanup
        print("Cleaning up temporary files...")
        for f_path in [temp_md_path, temp_intro_path, temp_content_unstyled_path]:
            if os.path.exists(f_path):
                try: os.remove(f_path)
                except OSError as rm_err: print(f"  Warning: Could not remove temporary file {f_path}: {rm_err}")

# --- UPDATED: Main Execution Logic ---
def create_remake(input_md_path, output_docx_path, template_path):
    """Creates a remake DOCX from Markdown via JSON intermediate."""
    print(f"Starting remake process for: {input_md_path}")
    try:
        # Read Original Text
        print("Reading original Markdown file...")
        with open(input_md_path, 'r', encoding='utf-8') as f:
            original_md_text = f.read()
        if not original_md_text.strip():
            print("Warning: Input Markdown file is empty.")
            return False # Cannot generate remake from empty file

        # Prepare base name
        base_name = os.path.splitext(os.path.basename(input_md_path))[0]
        if base_name.endswith("_extracted"): base_name = base_name[:-10]
        print(f"Using base name: {base_name}")

        # 1. Generate Initial JSON remake Content
        initial_remake_json = generate_remake_json_content(original_md_text)
        if initial_remake_json is None:
            print("ERROR: Failed to generate initial remake JSON from Gemini.")
            return False
        print("Initial Gemini JSON generation successful.")

        # 2. Verify and Correct the Initial JSON remake
        print("Verifying and correcting generated JSON remake...")
        verified_remake_json = verify_and_correct_remake_json(original_md_text, initial_remake_json)
        if verified_remake_json is None:
            print("ERROR: Failed during remake JSON verification/correction step.")
            return False
        print("Verification/correction of JSON step complete.")

        # 3. Convert Verified JSON to Markdown String
        print("Converting verified JSON to Markdown string...")
        final_markdown_content = json_to_markdown_remake(verified_remake_json)
        if not final_markdown_content:
             print("Warning: Generated Markdown content is empty.")
             # Optionally return False or continue to create an empty DOCX
             # return False
        print("Markdown string generated successfully.")
        # print("--- Generated Markdown ---") # Optional: Debug print
        # print(final_markdown_content)
        # print("------------------------")

        # 4. Process Markdown to Styled and Merged DOCX
        print("Starting DOCX processing (Pandoc, Styling, Merging)...")
        success = process_remake_md_to_docx(
            md_content=final_markdown_content, # Pass the generated MD
            lecture_name=base_name,
            template_path=template_path,
            output_docx_path=output_docx_path
        )

        if success:
            print(f"--- remake process completed successfully for: {input_md_path} ---")
        else:
            print(f"--- remake process failed for: {input_md_path} ---")

        return success

    except FileNotFoundError:
        print(f"ERROR: Input Markdown file '{input_md_path}' or template file '{template_path}' not found.")
        traceback.print_exc()
        return False
    except Exception as e:
        print(f"An UNEXPECTED error occurred during the remake generation process: {e}")
        traceback.print_exc()
        return False

# --- Example Usage (if needed for direct testing) ---
# if __name__ == "__main__":
#     # ... setup test paths from config ...
#     create_remake(test_md_path, test_docx_output, test_template)
