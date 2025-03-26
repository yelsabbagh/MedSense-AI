# summary_generator.py

import google.generativeai as genai
import config # Ensure config is imported
import os
import time
import json # <-- Added
import re   # <-- Added
from google.api_core.exceptions import ResourceExhausted
from docxtpl import DocxTemplate
import pypandoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
import lxml.etree as etree
from docx.oxml.shared import OxmlElement, qn
from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
import traceback # Import traceback for detailed error logging

# --- Helper Functions qn, _set_table_borders (Keep as before) ---
def qn(tag_name):
    """Stands for 'qualified name', a utility function for lxml."""
    import docx.oxml.ns
    prefix, tag = tag_name.split(':')
    uri = '{%s}%s' % (docx.oxml.ns.nsmap[prefix], tag)
    return uri

# --- Helper Function to set table borders ---
def _set_table_borders(table):
    """Sets single black 0.5pt borders for all cell edges in a table."""
    try:
        border_props = {
            "sz": "4", "val": "single", "color": "000000", "space": "0"
        }
        border_keys = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        tags = {key: f"w:{key}" for key in border_keys}

        tbl = table._tbl # Get the underlying CT_Tbl object

        # --- Correctly find or create tblPr ---
        tblPr = tbl.find(qn('w:tblPr')) # Find existing tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr') # Create if not found
            # Insert tblPr as the first child of tbl (common practice)
            tbl.insert(0, tblPr)
        # --- End of correction for tblPr ---

        # --- Correctly find or create tblBorders ---
        tblBorders = tblPr.find(qn('w:tblBorders')) # Find existing tblBorders within tblPr
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders') # Create if not found
            tblPr.append(tblBorders) # Append to tblPr
        # --- End of correction for tblBorders ---

        # Apply individual border settings (this part was okay)
        for key in border_keys:
            border_tag = tblBorders.find(qn(tags[key]))
            if border_tag is None:
                border_tag = OxmlElement(tags[key])
                tblBorders.append(border_tag)
            for prop, value in border_props.items():
                 border_tag.set(qn(f"w:{prop}"), value)

    except Exception as e:
        print(f"Error applying borders to table: {e}")
        raise # Re-raise the exception after printing

# --- Helper function _get_cell_text (Keep as before) ---
# def _get_cell_text(cell): ...

# --- Gemini API Call with Retry (MODIFIED FOR JSON) ---
def generate_with_retry(prompt, retries=5, delay=5, expect_json=False): # Added expect_json flag
    """Retries the generation request, handling JSON expectation."""
    genai.configure(api_key=config.API_KEY)

    gen_config = config.generation_config.copy()
    if expect_json:
        gen_config["response_mime_type"] = "application/json"
    else: # Ensure it's text/plain if not JSON
        gen_config["response_mime_type"] = "text/plain"

    model = genai.GenerativeModel(
        model_name=config.GEMINI_MODEL,
        generation_config=gen_config,
        safety_settings=config.safety_settings
    )

    for attempt in range(retries):
        try:
            print(f"Summary Gen: Sending request to Gemini (attempt {attempt + 1}, expecting {'JSON' if expect_json else 'Text'})...")
            response = model.generate_content(prompt)
            print("Summary Gen: Received response from Gemini.")
            # Simple check: if JSON was expected but we got blocked, return None early
            if expect_json:
                try:
                     # Check for blocking before accessing text
                     if response.prompt_feedback.block_reason:
                          print(f"Summary Gen: Content blocked due to: {response.prompt_feedback.block_reason}")
                          return None
                except Exception: pass # Ignore if feedback check fails
            return response.text # Return the text part (should be JSON string if expect_json)
        except ResourceExhausted as e:
            print(f"Summary Gen: Rate limit exceeded, retrying in {delay}s... ({attempt + 1}/{retries})")
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            print(f"Summary Gen: An unexpected error occurred during Gemini call: {e}")
            # ... (Error feedback logging) ...
            try: # Safely check feedback
                 if response and hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
                      print(f"Summary Gen: Content blocked due to: {response.prompt_feedback.block_reason}")
                 elif response and hasattr(response, 'candidates') and response.candidates and response.candidates[0].finish_reason != 'STOP':
                     print(f"Summary Gen: Generation stopped potentially due to: {response.candidates[0].finish_reason}")
            except Exception as feedback_err:
                 print(f"Summary Gen: Error accessing feedback details: {feedback_err}")
            # Retry logic
            if attempt < retries - 1:
                print(f"Summary Gen: Retrying after error in {delay}s...")
                time.sleep(delay)
                delay *= 2
            else:
                print("Summary Gen: Maximum retries reached or fatal error.")
                traceback.print_exc()
                if expect_json: return None # Return None on failure when expecting JSON
                raise e # Re-raise if not expecting JSON or want failure propagation

    print("Summary Gen: Max retries reached after ResourceExhausted errors.")
    return None # Return None if loop finishes without success

# --- UPDATED: Generate summary JSON Content using Gemini ---
def generate_summary_json_content(text):
    """Generates the initial structured summary as JSON using Gemini."""
    print("Summary Gen: Generating initial structured JSON via Gemini...")
    prompt = f"""
**Role:** You are an expert AI medical content specialist. Your task is to distill complex medical lecture text into a structured JSON format representing a concise summary suitable for review.

**Audience:** Medical students.

**Task:** Analyze the provided medical lecture text. Identify the most critical, high-yield information. Generate a JSON output representing this summary, organized into logical sections.

**JSON Output Structure:**
The output MUST be a single JSON list `[...]` containing section objects. Each section object MUST have the following structure:
```json
{{
  "title": "Section Title (e.g., Pathophysiology)", // String: Clear heading for the section
  "type": "paragraph | list | table",              // String: Specifies content format
  "content": "..."                                // Content format depends on 'type'
}}
```

**Content Formats based on "type":**
*   **If `type` is `"paragraph"`:** `content` MUST be a single string containing the summarized paragraph text. Use concise language.
*   **If `type` is `"list"`:** `content` MUST be a JSON array of strings `["item 1", "item 2", ...]`, where each string is a concise list item.
*   **If `type` is `"table"`:** `content` MUST be a JSON array of objects `[ {{...}}, {{...}} ]`. Each object represents a **row** in the table and MUST contain exactly two keys:
    *   `"key_point"`: String for the first column (e.g., term, category, feature).
    *   `"details"`: String for the second column (e.g., definition, description, value).
    Use this structure for comparisons, classifications, or key feature descriptions.

**Content Rules:**
1.  **High-Yield Focus:** Prioritize core concepts, mechanisms, clinical findings, diagnosis, treatment, etc.
2.  **Conciseness:** Keep all text (titles, paragraphs, list items, table cells) brief and to the point.
3.  **Accuracy:** MUST accurately reflect the source text. Do NOT add external information.
4.  **Logical Sections:** Group related information under appropriate `title` headings.

**Example JSON Output:**
```json
[
  {{
    "title": "Core Concept: Hypertension",
    "type": "paragraph",
    "content": "Defined as persistently elevated blood pressure (BP ≥ 130/80 mmHg based on newer guidelines). A major risk factor for cardiovascular disease."
  }},
  {{
    "title": "Types of Hypertension",
    "type": "table",
    "content": [
      {{ "key_point": "Primary (Essential)", "details": "No identifiable cause (~95% of cases). Multifactorial etiology (genetics, lifestyle)." }},
      {{ "key_point": "Secondary", "details": "Caused by an underlying condition (e.g., renal disease, endocrine disorders, medications)." }}
    ]
  }},
  {{
    "title": "Key Treatment Strategies",
    "type": "list",
    "content": [
      "Lifestyle Modifications (Diet, Exercise, Weight Loss)",
      "Pharmacological Therapy (Diuretics, ACEi, ARBs, CCBs, Beta-blockers)",
      "Address underlying cause if Secondary Hypertension"
    ]
  }}
]
```

**Input Text:**
---
{text}
---

**Final Output Instruction:**
Generate **ONLY** the JSON list `[...]` based on the requirements above. Do not include any other text, comments, or markdown formatting like ```json.
"""
    json_string = generate_with_retry(prompt, expect_json=True) # Expect JSON response
    if not json_string:
        print("Summary Gen: ERROR - Failed to get valid JSON response from Gemini for initial generation.")
        return None

    try:
        print("Summary Gen: Parsing Gemini JSON response...")
        # Clean potential markdown formatting
        if json_string.strip().startswith("```json"):
             json_string = json_string.strip()[7:-3].strip()
        elif json_string.strip().startswith("```"):
             json_string = json_string.strip()[3:-3].strip()
        summary_data = json.loads(json_string)
        # Basic validation: Check if it's a list
        if not isinstance(summary_data, list):
            print("Summary Gen: ERROR - Parsed JSON is not a list as expected.")
            print("--- Received Data ---")
            print(summary_data)
            print("---------------------")
            return None
        print("Summary Gen: Initial JSON structure parsed successfully.")
        return summary_data # Return the parsed list
    except json.JSONDecodeError as e:
        print(f"Summary Gen: ERROR - Failed to decode JSON response: {e}")
        print("--- Gemini Response Text (raw) ---")
        print(json_string)
        print("----------------------------------")
        return None
    except Exception as e:
        print(f"Summary Gen: An unexpected error occurred parsing JSON: {e}")
        traceback.print_exc()
        return None


# --- UPDATED: Verify and Correct Summary JSON ---
def verify_and_correct_summary_json(original_text, summary_json_to_verify):
    """
    Compares the generated summary JSON to the original_text for accuracy,
    completeness, conciseness, and adherence to JSON structure rules,
    refining it if necessary. Returns the *verified* JSON data (list).
    """
    print("Summary Gen: Starting JSON summary verification/correction...")

    # Convert the JSON object to a string for the prompt
    try:
        json_string_to_verify = json.dumps(summary_json_to_verify, indent=2)
    except Exception as e:
        print(f"Summary Gen: ERROR - Could not serialize JSON for verification prompt: {e}")
        return None # Cannot proceed if input JSON is invalid

    verification_rules = """
**Core Summary Verification Goals:**
1.  **Accuracy:** All information in the JSON summary MUST accurately reflect facts in the `ORIGINAL_TEXT`. Correct factual errors.
2.  **Key Point Coverage:** The JSON summary SHOULD capture the most important, high-yield points from the `ORIGINAL_TEXT`. Add critical missing points concisely *if structure allows*. Remove trivial details.
3.  **Conciseness:** All strings within the JSON (`title`, `content` values) MUST be brief and avoid redundancy.
4.  **No External Information:** Ensure NO information is present in the JSON that is not derivable from the `ORIGINAL_TEXT`. Remove any external info.
5.  **JSON Structure Adherence:** The output MUST be a valid JSON list `[...]`. Each object within the list MUST contain `title` (string), `type` (string: "paragraph", "list", or "table"), and `content`. The `content` format MUST match the `type` as specified (string for paragraph, list of strings for list, list of {"key_point": string, "details": string} objects for table). Correct any structural errors.
"""

    prompt = f"""
**Role:** You are an expert AI medical content editor specializing in verifying and refining structured JSON summaries.

**Objective:** Analyze the `JSON_SUMMARY_TO_VERIFY`. Compare it against the `ORIGINAL_TEXT` based on the `VERIFICATION_GOALS`. Ensure the JSON is accurate, concise, covers key points, contains no external info, and strictly adheres to the specified JSON structure.

{verification_rules}

---
**Input 1: ORIGINAL_TEXT** (The full source material)

{original_text}

---
**Input 2: JSON_SUMMARY_TO_VERIFY** (The JSON generated previously)
```json
{json_string_to_verify}
```

---
**Verification & Refinement Task:**
Review the `JSON_SUMMARY_TO_VERIFY` section by section against the `ORIGINAL_TEXT`.
1.  **Check Accuracy:** Correct any factual errors in `title` or `content` fields.
2.  **Check Key Points:** Add/remove information to ensure high-yield coverage without excessive detail.
3.  **Check Conciseness:** Shorten verbose text within the JSON strings.
4.  **Check External Info:** Remove any data not found in `ORIGINAL_TEXT`.
5.  **Check JSON Structure:** Ensure the entire output is a valid JSON list `[...]` and each section object conforms *exactly* to the structure (`title`, `type`, `content` format matching `type`). Fix any structural mistakes.

---
**Output Instruction:**
Output **ONLY** the final, verified, and potentially refined JSON list `[...]`.
*   If refinements were made, output the improved JSON list.
*   If the original `JSON_SUMMARY_TO_VERIFY` met all goals perfectly, output it exactly as provided.
*   **Do NOT** include any explanations, comments, confirmations, or conversational text. Your entire response must be the final JSON list content, starting with `[` and ending with `]`.
"""

    verified_json_string = generate_with_retry(prompt, expect_json=True) # Expect JSON back
    if not verified_json_string:
        print("Summary Gen: ERROR - Failed to get valid JSON response from Gemini for verification.")
        return None

    try:
        print("Summary Gen: Parsing verified Gemini JSON response...")
        # Clean potential markdown formatting
        if verified_json_string.strip().startswith("```json"):
             verified_json_string = verified_json_string.strip()[7:-3].strip()
        elif verified_json_string.strip().startswith("```"):
             verified_json_string = verified_json_string.strip()[3:-3].strip()
        verified_summary_data = json.loads(verified_json_string)
        # Basic validation
        if not isinstance(verified_summary_data, list):
            print("Summary Gen: ERROR - Verified JSON is not a list as expected.")
            print("--- Received Data ---")
            print(verified_summary_data)
            print("---------------------")
            return None
        print("Summary Gen: Verified JSON structure parsed successfully.")
        return verified_summary_data # Return the parsed verified list
    except json.JSONDecodeError as e:
        print(f"Summary Gen: ERROR - Failed to decode verified JSON response: {e}")
        print("--- Verified Gemini Response Text (raw) ---")
        print(verified_json_string)
        print("-----------------------------------------")
        return None
    except Exception as e:
        print(f"Summary Gen: An unexpected error occurred parsing verified JSON: {e}")
        traceback.print_exc()
        return None

# --- NEW: Function to Convert JSON Summary to Markdown ---
def json_to_markdown_summary(summary_data):
    """Converts the structured summary data (list of sections) into a Markdown string."""
    if not summary_data:
        print("Summary Gen: Warning - JSON summary data is empty, generating empty Markdown.")
        return ""

    md_parts = []
    for section in summary_data:
        title = section.get("title", "Untitled Section")
        type = section.get("type")
        content = section.get("content")

        # Add Heading
        md_parts.append(f"## {title}\n") # Use H2 for sections

        if type == "paragraph" and isinstance(content, str):
            md_parts.append(f"{content}\n")
        elif type == "list" and isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    md_parts.append(f"- {item}\n")
            md_parts.append("\n") # Add space after list
        elif type == "table" and isinstance(content, list) and content:
            # Check if content is a list of dicts with expected keys
            if all(isinstance(row, dict) and "key_point" in row and "details" in row for row in content):
                # Create Markdown Table
                md_parts.append("| Key Point | Details |\n") # Fixed headers
                md_parts.append("|---|---|\n")
                for row in content:
                    # Escape pipe characters within cell content
                    key_point_cell = str(row.get("key_point", "")).replace('|', '\\|')
                    details_cell = str(row.get("details", "")).replace('|', '\\|')
                    md_parts.append(f"| {key_point_cell} | {details_cell} |\n")
                md_parts.append("\n") # Add space after table
            else:
                print(f"Summary Gen: Warning - Section '{title}' has type 'table' but content format is invalid. Skipping table generation.")
                md_parts.append(f"[Content for section '{title}' intended as table, but format was invalid]\n")
        else:
            print(f"Summary Gen: Warning - Section '{title}' has unknown type '{type}' or invalid content. Skipping content.")
            md_parts.append(f"[Content for section '{title}' has unknown type '{type}' or invalid content]\n")

        md_parts.append("\n") # Add space between sections

    return "".join(md_parts)


# --- UPDATED: Function to Apply Styling (Removed Row Merging) ---
def apply_styling_to_summary_docx(docx_path):
    """Loads a DOCX file, applies table borders, paragraph styling, and saves in place."""
    print(f"Summary Styling: Applying styling to: {docx_path}")
    try:
        doc = Document(docx_path)

        # --- Style Tables (Borders and Font) ---
        print(f"  Found {len(doc.tables)} tables.")
        for table_idx, table in enumerate(doc.tables):
            try:
                # --- Apply Borders ---
                _set_table_borders(table)

                # --- Apply Font & Bold First Column ---
                # Pandoc MD tables usually have a header row, skip it for bolding data
                has_header = (len(table.rows) > 1 and
                                table.rows[0].cells and len(table.rows[0].cells) > 0 and
                                table.rows[0].cells[0].text.strip().lower() == "key point") # Check header based on our generation

                for row_idx, row in enumerate(table.rows):
                    # Bold first column data cells (skip header)
                    if has_header and row_idx == 0:
                         pass # Skip actual header row if detected
                    elif len(row.cells) > 0:
                         for paragraph in row.cells[0].paragraphs:
                               for run in paragraph.runs:
                                   run.bold = True

                    # Apply font to all cells
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Poppins'
                                # Apply Oxml font settings (optional but good practice)
                                r = run._element; rPr = r.get_or_add_rPr()
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
                                rFonts.set(qn('w:ascii'), 'Poppins'); rFonts.set(qn('w:hAnsi'), 'Poppins')
                                rFonts.set(qn('w:eastAsia'), 'Poppins'); rFonts.set(qn('w:cs'), 'Poppins')

            except Exception as table_style_error:
                print(f"  ERROR: Failed styling table {table_idx+1}: {table_style_error}")
                traceback.print_exc()

        # --- Apply Other Formatting (non-table paragraphs) ---
        for paragraph in doc.paragraphs:
             # Font (Apply Poppins)
             for run in paragraph.runs:
                run.font.name = 'Poppins'
                r = run._element; rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Poppins'); rFonts.set(qn('w:hAnsi'), 'Poppins')
                rFonts.set(qn('w:eastAsia'), 'Poppins'); rFonts.set(qn('w:cs'), 'Poppins')

             # Alignment, spacing, list indents (Keep existing logic)
             paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
             paragraph.paragraph_format.line_spacing = 1.0
             if paragraph.style.name.startswith('Heading'):
                 paragraph.paragraph_format.left_indent = Inches(0)
                 paragraph.paragraph_format.space_before = Pt(12)
                 paragraph.paragraph_format.space_after = Pt(6)
             elif paragraph.style.name.startswith('List'):
                  # Let Pandoc handle list indentation, just ensure alignment/direction
                  pass # Or apply minimal indent if needed: paragraph.paragraph_format.left_indent = Inches(0.25)
             else: # Regular paragraphs
                  paragraph.paragraph_format.left_indent = Inches(0) # Often 0 indent is fine for summaries
                  paragraph.paragraph_format.first_line_indent = Inches(0)

             # LTR direction (Keep existing logic)
             p = paragraph._element; pPr = p.get_or_add_pPr()
             bidi_tag = pPr.find(qn('w:bidi'))
             if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
             bidi_tag.set(qn('w:val'), "0")

        # Save the styled document back to the same path
        doc.save(docx_path)
        print(f"  Styling applied successfully to {docx_path}")
        return True

    except Exception as e:
        print(f"ERROR applying styling to {docx_path}: {e}")
        traceback.print_exc()
        return False


# --- Merge Template and Content (Keep as before) ---
def merge_template_and_styled_content(intro_docx_path, styled_content_path, final_output_path):
    # ... (No changes needed here) ...
    print(f"Merging '{intro_docx_path}' and '{styled_content_path}' into '{final_output_path}'")
    try:
        intro_doc = Document(intro_docx_path)
        intro_doc.add_page_break()
        content_doc = Document(styled_content_path)
        for element in content_doc.element.body:
            intro_doc.element.body.append(element)

        # Final LTR/Alignment Pass
        print("  Applying final LTR/Alignment pass to merged document...")
        for paragraph in intro_doc.paragraphs:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = paragraph._element; pPr = p.get_or_add_pPr()
            bidi_tag = pPr.find(qn('w:bidi'))
            if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
            bidi_tag.set(qn('w:val'), "0")
        for table in intro_doc.tables:
             for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p = paragraph._element; pPr = p.get_or_add_pPr()
                        bidi_tag = pPr.find(qn('w:bidi'))
                        if bidi_tag is None: bidi_tag = etree.SubElement(pPr, qn('w:bidi'))
                        bidi_tag.set(qn('w:val'), "0")

        intro_doc.save(final_output_path)
        print(f"  Successfully merged and saved to {final_output_path}")
        return True
    except Exception as e:
        print(f"ERROR during merging/final saving: {e}")
        traceback.print_exc()
        return False


# --- Conversion and Styling Orchestration (Keep as before) ---
def process_summary_md_to_docx(md_content, lecture_name, template_path, output_docx_path):
    """Orchestrates MD save, Pandoc conversion, styling, and merging."""
    # ... (No changes needed here - it takes final Markdown string) ...
    temp_md_path = f"temp_summary_{lecture_name}.md" # Unique temp name
    temp_intro_path = f"temp_summary_intro_{lecture_name}.docx"
    temp_content_unstyled_path = f"temp_summary_content_unstyled_{lecture_name}.docx"
    try:
        # 1. Render template intro
        print("Loading and rendering template...")
        doc_template = DocxTemplate(template_path)
        context = {'lecture_name': lecture_name}
        doc_template.render(context)
        doc_template.save(temp_intro_path)
        # 2. Save generated Markdown content
        print(f"Saving generated Markdown to '{temp_md_path}'...")
        with open(temp_md_path, 'w', encoding='utf-8') as f_md: f_md.write(md_content)
        # 3. Convert MD to DOCX via Pandoc
        print(f"Converting MD to DOCX via Pandoc: '{temp_md_path}' -> '{temp_content_unstyled_path}'...")
        pypandoc_args = ['--wrap=none']
        pandoc_path = getattr(config, 'PANDOC_PATH', None)
        if pandoc_path and os.path.exists(pandoc_path):
            try:
                 print(f"  Using custom Pandoc path: {pandoc_path}")
                 pypandoc_args.extend(['--pandoc-binary', pandoc_path])
            except Exception as pe: print(f"  Warning: Issue using custom Pandoc path {pandoc_path}: {pe}.")
        else:
             if pandoc_path: print(f"  Warning: Custom Pandoc path not found: {pandoc_path}. Relying on PATH.")
        pypandoc.convert_file(temp_md_path, 'docx', outputfile=temp_content_unstyled_path, extra_args=pypandoc_args)
        print("  Pandoc conversion successful.")
        # 4. Apply styling (UPDATED function)
        styling_success = apply_styling_to_summary_docx(temp_content_unstyled_path)
        if not styling_success: print("ERROR: Failed to apply styling."); return False
        # 5. Merge
        merge_success = merge_template_and_styled_content(temp_intro_path, temp_content_unstyled_path, output_docx_path)
        if not merge_success: print("ERROR: Failed to merge template and styled content."); return False
        print(f"Summary DOCX generated successfully: {output_docx_path}")
        return True
    except pypandoc.PandocMissingError: print("ERROR: Pandoc executable not found."); traceback.print_exc(); return False
    except FileNotFoundError as fnf_error: print(f"ERROR: File not found: {fnf_error}"); traceback.print_exc(); return False
    except Exception as e: print(f"An CRITICAL error occurred during DOCX orchestration: {e}"); traceback.print_exc(); return False
    finally: # Cleanup
        print("Cleaning up temporary files...")
        for f_path in [temp_md_path, temp_intro_path, temp_content_unstyled_path]:
            if os.path.exists(f_path):
                try: os.remove(f_path)
                except OSError as rm_err: print(f"  Warning: Could not remove temporary file {f_path}: {rm_err}")


# --- UPDATED: Main Execution Logic ---
def create_summary(input_md_path, output_docx_path, template_path):
    """Creates a summary DOCX from Markdown via JSON intermediate."""
    print(f"Starting summary process for: {input_md_path}")
    try:
        # Read Original Text
        print("Reading original Markdown file...")
        with open(input_md_path, 'r', encoding='utf-8') as f:
            original_md_text = f.read()
        if not original_md_text.strip():
            print("Warning: Input Markdown file is empty.")
            return False # Cannot generate summary from empty file

        # Prepare base name
        base_name = os.path.splitext(os.path.basename(input_md_path))[0]
        if base_name.endswith("_extracted"): base_name = base_name[:-10]
        print(f"Using base name: {base_name}")

        # 1. Generate Initial JSON Summary Content
        initial_summary_json = generate_summary_json_content(original_md_text)
        if initial_summary_json is None:
            print("ERROR: Failed to generate initial summary JSON from Gemini.")
            return False
        print("Initial Gemini JSON generation successful.")

        # 2. Verify and Correct the Initial JSON Summary
        print("Verifying and correcting generated JSON summary...")
        verified_summary_json = verify_and_correct_summary_json(original_md_text, initial_summary_json)
        if verified_summary_json is None:
            print("ERROR: Failed during summary JSON verification/correction step.")
            return False
        print("Verification/correction of JSON step complete.")

        # 3. Convert Verified JSON to Markdown String
        print("Converting verified JSON to Markdown string...")
        final_markdown_content = json_to_markdown_summary(verified_summary_json)
        if not final_markdown_content:
             print("Warning: Generated Markdown content is empty.")
             # Optionally return False or continue to create an empty DOCX
             # return False
        print("Markdown string generated successfully.")
        # print("--- Generated Markdown ---") # Optional: Debug print
        # print(final_markdown_content)
        # print("------------------------")

        # 4. Process Markdown to Styled and Merged DOCX
        print("Starting DOCX processing (Pandoc, Styling, Merging)...")
        success = process_summary_md_to_docx(
            md_content=final_markdown_content, # Pass the generated MD
            lecture_name=base_name,
            template_path=template_path,
            output_docx_path=output_docx_path
        )

        if success:
            print(f"--- summary process completed successfully for: {input_md_path} ---")
        else:
            print(f"--- summary process failed for: {input_md_path} ---")

        return success

    except FileNotFoundError:
        print(f"ERROR: Input Markdown file '{input_md_path}' or template file '{template_path}' not found.")
        traceback.print_exc()
        return False
    except Exception as e:
        print(f"An UNEXPECTED error occurred during the summary generation process: {e}")
        traceback.print_exc()
        return False

# --- Example Usage (if needed for direct testing) ---
# if __name__ == "__main__":
#     # ... setup test paths from config ...
#     create_summary(test_md_path, test_docx_output, test_template)
